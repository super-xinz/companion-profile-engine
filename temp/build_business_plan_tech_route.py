from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "商业计划书_技术路线_一页成稿_20260807.docx"

# Design authority: standard_business_brief.
# Named override "investor_one_page": compact 9.2 pt body and restrained navy/teal
# palette for a dense, single-page business-plan insert. Page geometry remains the
# preset's US Letter portrait with 1-inch margins and 9360 DXA usable width.

FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"
NAVY = "183B56"
BLUE = "2E74B5"
TEAL = "0F8B8D"
SKY = "EAF3F8"
PALE_TEAL = "E8F5F3"
PALE_GOLD = "FFF5DB"
PALE_GRAY = "F4F6F8"
MID_GRAY = "64748B"
DARK = "172B3A"
WHITE = "FFFFFF"
LINE = "D9E2E8"
GOLD = "DFAF2B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = qn(f"w:{name}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, spec in edges.items():
        tag = qn(f"w:{edge_name}")
        edge = borders.find(tag)
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key in ("val", "sz", "space", "color"):
            if key in spec:
                edge.set(qn(f"w:{key}"), str(spec[key]))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def set_run(run, *, size: float, color: str = DARK, bold=False, italic=False) -> None:
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph(
    paragraph,
    *,
    before=0,
    after=0,
    line=1.0,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    keep_next=False,
) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next
    paragraph.alignment = align


def clear_cell(cell) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph(p)


def add_text(cell, text: str, *, size: float, color=DARK, bold=False, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[-1] if not cell.paragraphs[-1].text else cell.add_paragraph()
    set_paragraph(p, after=after, line=line, align=align)
    run = p.add_run(text)
    set_run(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    set_paragraph(p, before=6, after=4, line=1.0, keep_next=True)
    run = p.add_run(text)
    set_run(run, size=13, color=NAVY, bold=True)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    set_paragraph(hp, after=0, line=1.0)
    left = hp.add_run("商业计划书 · 技术路线")
    set_run(left, size=8.3, color=MID_GRAY, bold=True)
    tab_stops = hp.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5))
    right = hp.add_run("\t共鸣画像引擎 × HIWM")
    set_run(right, size=8.3, color=MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    set_paragraph(fp, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = fp.add_run("TECHNOLOGY ROADMAP · 2026")
    set_run(run, size=7.5, color=MID_GRAY, bold=True)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    add_header_footer(doc)

    # memo_masthead-derived title stack (without a separate cover or metadata block).
    kicker = doc.add_paragraph()
    set_paragraph(kicker, after=2, line=1.0)
    set_run(kicker.add_run("BUSINESS PLAN · TECHNOLOGY ROADMAP"), size=8.7, color=TEAL, bold=True)

    title = doc.add_paragraph()
    set_paragraph(title, after=3, line=1.0, keep_next=True)
    set_run(title.add_run("从“生成回答”到“预测互动后果”"), size=21.5, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph(subtitle, after=8, line=1.0, keep_next=True)
    set_run(
        subtitle.add_run("共鸣画像引擎 × HIWM：面向陪伴机器人与 B 端智能体的可进化互动决策中枢"),
        size=10.5,
        color=MID_GRAY,
        bold=True,
    )

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    clear_cell(cell)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=135, start=120, bottom=135, end=120)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 4, "color": NAVY},
        bottom={"val": "single", "sz": 4, "color": NAVY},
        start={"val": "single", "sz": 4, "color": NAVY},
        end={"val": "single", "sz": 4, "color": NAVY},
    )
    p = cell.paragraphs[0]
    set_paragraph(p, line=1.08)
    set_run(p.add_run("核心路径｜"), size=10.2, color="8FE0D6", bold=True)
    set_run(
        p.add_run("持续理解同一个人 → 开口前推演三种沟通策略 → 选择风险更低、信息增益更高的一条 → 用真实回应持续纠偏。"),
        size=10.2,
        color=WHITE,
        bold=True,
    )
    p2 = cell.add_paragraph()
    set_paragraph(p2, before=2, line=1.05)
    set_run(
        p2.add_run("当前已完成可运行 MVP 与双服务集成；下一阶段的技术投入重点是数据校准、垂直场景验证与规模化部署，而非从零搭建 Demo。"),
        size=8.6,
        color="DDEAF2",
    )

    add_heading(doc, "核心技术闭环｜每一轮互动都沉淀为可验证的数据资产")

    steps = [
        ("01", "有界多模态感知", "final ASR + 语音韵律/可观察行为 + 用户确认上下文；默认不保存原始音视频。", SKY, BLUE),
        ("02", "证据化动态画像", "17 维画像、MBTI 连续维度、长期记忆与短期状态；生日数码仅作低置信先验，九型仅接受明确授权输入。", PALE_TEAL, TEAL),
        ("03", "行动条件化推演", "维护任务相关互动状态；每轮生成 3 条候选动作，预测下一反应、状态变化与未知项。", PALE_GOLD, GOLD),
        ("04", "风险约束规划", "按目标、信息增益、风险与不确定性联合评分；证据不足或风险过高时澄清、等待或转人工。", PALE_GOLD, GOLD),
        ("05", "可验证执行", "执行前以 JSONL + fsync + SHA-256 锁定预测；TTS/机器人只读语音命令使用 turn_id 幂等。", SKY, BLUE),
        ("06", "反馈闭环学习", "用下一轮真实回应评估匹配与偏差，更新工作信念和画像版本，形成“证据—动作—结果”轨迹。", PALE_TEAL, TEAL),
    ]
    grid = doc.add_table(rows=2, cols=3)
    set_table_geometry(grid, [3120, 3120, 3120])
    for idx, (num, heading, body, fill, accent) in enumerate(steps):
        cell = grid.cell(idx // 3, idx % 3)
        clear_cell(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=115, start=120, bottom=110, end=120)
        white_border = {"val": "single", "sz": 8, "color": WHITE}
        set_cell_border(
            cell,
            top={"val": "single", "sz": 18, "color": accent},
            bottom=white_border,
            start=white_border,
            end=white_border,
        )
        p = cell.paragraphs[0]
        set_paragraph(p, after=3, line=1.0, keep_next=True)
        set_run(p.add_run(f"{num}  "), size=8.3, color=accent, bold=True)
        set_run(p.add_run(heading), size=9.2, color=NAVY, bold=True)
        body_p = cell.add_paragraph()
        set_paragraph(body_p, line=1.08)
        set_run(body_p.add_run(body), size=7.85, color=DARK)

    add_heading(doc, "工程基础与升级路线｜从可部署 MVP 走向可校准、可复制的行业能力")

    phases = [
        (
            "已完成｜可部署 MVP",
            "Vue/FastAPI 实时交互 + 独立画像状态机 + BFF + PostgreSQL；多租户鉴权、幂等、乐观锁、规则审批、知情同意与审计链已工程化。",
            SKY,
            BLUE,
        ),
        (
            "0–6 个月｜场景化验证",
            "围绕陪伴、学习/心理协作及机器人客户建立试点；统一状态字段与结果标签，采集经同意的完整轨迹，补齐机器人回执、失败恢复和私有化部署。",
            PALE_TEAL,
            TEAL,
        ),
        (
            "6–18 个月｜校准与规模化",
            "以 NLL、Brier、ECE、行动排序和多模态消融验证增量价值；迭代多步轨迹模型与受控在线更新，扩展 ROS/厂商 SDK 和垂直行业策略包。",
            PALE_GOLD,
            GOLD,
        ),
    ]
    road = doc.add_table(rows=2, cols=3)
    set_table_geometry(road, [3120, 3120, 3120])
    set_repeat_table_header(road.rows[0])
    for col, (label, body, fill, accent) in enumerate(phases):
        hcell = road.cell(0, col)
        clear_cell(hcell)
        hcell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hcell, accent)
        set_cell_margins(hcell, top=80, start=120, bottom=80, end=120)
        set_cell_border(
            hcell,
            top={"val": "single", "sz": 6, "color": WHITE},
            bottom={"val": "single", "sz": 6, "color": WHITE},
            start={"val": "single", "sz": 6, "color": WHITE},
            end={"val": "single", "sz": 6, "color": WHITE},
        )
        hp = hcell.paragraphs[0]
        set_paragraph(hp, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
        set_run(hp.add_run(label), size=8.25, color=WHITE, bold=True)

        bcell = road.cell(1, col)
        clear_cell(bcell)
        bcell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(bcell, fill)
        set_cell_margins(bcell, top=105, start=120, bottom=105, end=120)
        set_cell_border(
            bcell,
            top={"val": "single", "sz": 6, "color": WHITE},
            bottom={"val": "single", "sz": 6, "color": WHITE},
            start={"val": "single", "sz": 6, "color": WHITE},
            end={"val": "single", "sz": 6, "color": WHITE},
        )
        bp = bcell.paragraphs[0]
        set_paragraph(bp, line=1.08)
        set_run(bp.add_run(body), size=7.75, color=DARK)

    add_heading(doc, "可形成的技术壁垒")
    moat = doc.add_table(rows=1, cols=4)
    set_table_geometry(moat, [2340, 2340, 2340, 2340])
    moat_items = [
        ("数据壁垒", "经同意的个体互动轨迹，而非通用语料"),
        ("算法壁垒", "个体状态 + 行动条件化预测 + 概率校准"),
        ("工程壁垒", "证据/版本/审计 + 执行前预测锁 + 安全退出"),
        ("生态壁垒", "可替换模型栈 + B 端 API + 机器人适配层"),
    ]
    for col, (label, detail) in enumerate(moat_items):
        cell = moat.cell(0, col)
        clear_cell(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, PALE_GRAY if col % 2 == 0 else SKY)
        set_cell_margins(cell, top=95, start=120, bottom=95, end=120)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 6, "color": WHITE},
            bottom={"val": "single", "sz": 6, "color": WHITE},
            start={"val": "single", "sz": 6, "color": WHITE},
            end={"val": "single", "sz": 6, "color": WHITE},
        )
        lp = cell.paragraphs[0]
        set_paragraph(lp, after=2, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
        set_run(lp.add_run(label), size=8.1, color=NAVY, bold=True)
        dp = cell.add_paragraph()
        set_paragraph(dp, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.04)
        set_run(dp.add_run(detail), size=7.25, color=MID_GRAY)

    closing = doc.add_paragraph()
    set_paragraph(closing, before=4, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(closing.add_run("投资逻辑：用资本加速“真实场景数据 → 概率校准 → 行业复制”的飞轮，把工程先发优势转化为持续提升的互动决策模型。"), size=8.55, color=TEAL, bold=True)

    # Document metadata.
    doc.core_properties.title = "商业计划书技术路线｜共鸣画像引擎 × HIWM"
    doc.core_properties.subject = "可直接纳入商业计划书的一页技术路线"
    doc.core_properties.author = "项目团队"
    doc.core_properties.keywords = "HIWM, 共鸣画像引擎, 技术路线, 商业计划书"
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
