from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(r"D:\Users\Lenovo\Desktop\companion-profile-engine")
CREDENTIAL_FILE = ROOT / "robot-company-live-credentials-20260807.txt"
OUTPUT = ROOT / "机器人公司_API完整接入指南_20260807.docx"
TABLE_HELPER_DIR = Path(
    r"C:\Users\Lenovo\.codex\plugins\cache\openai-primary-runtime\documents\26.805.11740\skills\documents\scripts"
)
sys.path.insert(0, str(TABLE_HELPER_DIR))
from table_geometry import apply_table_geometry  # noqa: E402


BASE_URL = "https://companion-profile-engine.zeabur.app"
DOCS_URL = BASE_URL + "/docs"
OPENAPI_URL = BASE_URL + "/openapi.json"

COLORS = {
    "navy": "17365D",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "0B2545",
    "muted": "5F6B7A",
    "light_blue": "E8EEF5",
    "light_gray": "F2F4F7",
    "callout": "F4F6F9",
    "warning": "FFF4CE",
    "warning_text": "7A5A00",
    "success": "E6F4EA",
    "success_text": "1E6B3A",
    "red": "9B1C1C",
    "white": "FFFFFF",
    "border": "CDD5DF",
}


def read_credentials() -> tuple[str, str]:
    values: dict[str, str] = {}
    for line in CREDENTIAL_FILE.read_text(encoding="utf-8").splitlines():
        if "=" in line:
            key, value = line.split("=", 1)
            values[key.strip()] = value.strip()
    tenant = values.get("TENANT_ID", "")
    key = values.get("API_KEY", "")
    if not tenant or not re.fullmatch(r"[0-9a-f]{64}", key):
        raise RuntimeError("生产凭据文件缺失或格式错误")
    return tenant, key


TENANT_ID, API_KEY = read_credentials()


def set_run_font(run, *, name="Calibri", east_asia="Microsoft YaHei", size=11,
                 color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, color="CDD5DF", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("第 ")
    set_run_font(prefix, size=9, color=COLORS["muted"])
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=COLORS["muted"])


def set_keep_with_next(paragraph, keep=True):
    paragraph.paragraph_format.keep_with_next = keep


def add_body(doc, text="", *, bold_prefix=None, after=6, color="000000"):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True, color=color)
        set_run_font(p.add_run(text[len(bold_prefix):]), color=color)
    else:
        set_run_font(p.add_run(text), color=color)
    return p


def add_kv_paragraph(doc, label, value, *, value_color="000000", after=4):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    set_run_font(p.add_run(label + "："), bold=True, color=COLORS["ink"])
    set_run_font(p.add_run(value), color=value_color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level],
                 bold=True, color={1: COLORS["blue"], 2: COLORS["blue"], 3: COLORS["dark_blue"]}[level])
    return p


def add_callout(doc, label, text, *, fill="F4F6F9", label_color="1F3A5F"):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), label_color)
    borders.append(left)
    p_pr.append(borders)
    set_run_font(p.add_run(label + "  "), bold=True, color=label_color)
    set_run_font(p.add_run(text), color="263238")
    return p


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.06)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), COLORS["border"])
        borders.append(el)
    p_pr.append(borders)
    run = p.add_run(code)
    set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=8.2, color="24292F")
    return p


def add_bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="Guide Bullet")
    p.paragraph_format.left_indent = Twips(540 + level * 360)
    p.paragraph_format.first_line_indent = Twips(-270)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(BULLET_NUM_ID))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    set_run_font(p.add_run(text))
    return p


def add_number(doc, text, *, level=0):
    global NUMBER_NUM_ID
    if not doc.paragraphs or doc.paragraphs[-1].style.name != "Guide Number":
        NUMBER_NUM_ID = add_num_instance(doc, NUMBER_ABSTRACT_ID)
    p = doc.add_paragraph(style="Guide Number")
    p.paragraph_format.left_indent = Twips(540 + level * 360)
    p.paragraph_format.first_line_indent = Twips(-270)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(NUMBER_NUM_ID))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    set_run_font(p.add_run(text))
    return p


def add_num_instance(doc, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    num_id = max(num_ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_numbering(doc, fmt, text) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text if fmt == "bullet" else f"%{level + 1}.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)
    # OOXML requires every abstractNum to precede all num instances.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)
    num_id = add_num_instance(doc, abstract_id)
    return num_id, abstract_id


def add_table(doc, headers, rows, widths, *, header_fill="E8EEF5", aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0].height = None
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(cell, header_fill)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(str(header)), size=9.2, bold=True, color=COLORS["ink"])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if aligns and aligns[idx] == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(str(value)), size=9.2)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=120,
                         cell_margins_dxa={"top": 100, "bottom": 100, "start": 120, "end": 120})
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    return table


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    heading_tokens = {
        1: (16, COLORS["blue"], 18, 10),
        2: (13, COLORS["blue"], 14, 7),
        3: (12, COLORS["dark_blue"], 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("Guide Bullet", "Guide Number"):
        style = doc.styles.add_style(name, 1)
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("Companion Profile Engine  |  机器人公司生产 API 接入"),
                 size=8.5, color=COLORS["muted"])
    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_page_number(p)
    section.different_first_page_header_footer = False


doc = Document()
configure_page(doc)
configure_styles(doc)
BULLET_NUM_ID, BULLET_ABSTRACT_ID = add_numbering(doc, "bullet", "•")
NUMBER_NUM_ID, NUMBER_ABSTRACT_ID = add_numbering(doc, "decimal", "%1.")

props = doc.core_properties
props.title = "机器人公司 API 完整接入指南"
props.subject = "Companion Profile Engine v0.2.0 生产接口"
props.author = "Companion Profile Engine 项目组"
props.keywords = "API, 机器人, 用户画像, 接入指南"

# Opening block: customer_pack pattern, compact operational variant.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(2)
set_run_font(p.add_run("PRODUCTION INTEGRATION GUIDE"), size=9.5, bold=True, color=COLORS["blue"])
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
set_run_font(p.add_run("Companion Profile Engine"), size=27, bold=True, color=COLORS["navy"])
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
set_run_font(p.add_run("机器人公司生产 API 完整接入指南"), size=15, color=COLORS["dark_blue"])

add_table(
    doc,
    ["交付项", "内容"],
    [
        ("接口版本", "v0.2.0"),
        ("指南日期", "2026-08-07"),
        ("运行状态", "已上线；HTTPS、应用和 PostgreSQL 正常"),
        ("接入模式", "机器人公司服务端通过 HTTPS 调用，无需部署我方源码"),
    ],
    [2700, 6660],
)
add_callout(doc, "交付结论", "该生产 API 已完成真实线上验收（10/10 通过）。负责人只需按本文配置服务端凭据并依次完成健康检查、用户初始化和消息摄取，即可投入联调。", fill=COLORS["success"], label_color=COLORS["success_text"])
add_callout(doc, "机密提醒", "本文含真实生产 API Key。只能发给机器人公司负责接入的服务端负责人；请勿发群、勿放前端、勿提交代码仓库、勿写入日志。若误发，请立即联系交付方轮换密钥。", fill=COLORS["warning"], label_color=COLORS["warning_text"])

add_heading(doc, "1. 负责人实际需要做什么", 1)
add_number(doc, "在机器人公司后端的密钥管理系统中保存 BASE_URL、TENANT_ID、API_KEY。")
add_number(doc, "确认后端服务器可出站访问 companion-profile-engine.zeabur.app 的 HTTPS 443 端口；不需要开放入站端口。")
add_number(doc, "按第 3 节执行健康检查和鉴权检查。")
add_number(doc, "按第 5 节实现“读取/初始化 → 摄取消息 → 使用 reply_hints 生成回复”的主链路。")
add_number(doc, "按第 12 节完成上线验收。")
add_body(doc, "除本文外，不需要额外安装包、源码、数据库账号或云平台账号。OpenAPI 文档可以在线查看，但正式集成以本文约定为准。")

add_heading(doc, "2. 生产环境地址与正式凭据", 1)
add_table(
    doc,
    ["变量", "正式值"],
    [
        ("BASE_URL", BASE_URL),
        ("TENANT_ID", TENANT_ID),
        ("API_KEY", API_KEY),
        ("Swagger UI", DOCS_URL),
        ("OpenAPI JSON", OPENAPI_URL),
    ],
    [2200, 7160],
)
add_body(doc, "建议配置为后端环境变量：")
add_code(doc, f"PROFILE_ENGINE_BASE_URL={BASE_URL}\nPROFILE_ENGINE_TENANT_ID={TENANT_ID}\nPROFILE_ENGINE_API_KEY={API_KEY}")
add_callout(doc, "凭据规则", "API Key 与租户 ID 是一组固定凭据。不得只凭 API Key 允许调用任意租户；所有请求必须同时发送这两个 Header。", fill=COLORS["callout"], label_color=COLORS["dark_blue"])

add_heading(doc, "3. 五分钟连通性验证", 1)
add_heading(doc, "3.1 健康检查（无需鉴权）", 2)
add_code(doc, f'curl.exe -sS "{BASE_URL}/health"')
add_body(doc, "必须得到 HTTP 200，且 JSON 同时满足 status=ok、services.application=ok、services.database=ok：")
add_code(doc, '{\n  "status": "ok",\n  "service": "companion-profile-engine",\n  "version": "0.2.0",\n  "services": {"application": "ok", "database": "ok"}\n}')
add_heading(doc, "3.2 鉴权与画像读取", 2)
add_code(doc, f'curl.exe -sS "{BASE_URL}/v1/profiles/connectivity-check" ^\n  -H "X-Tenant-ID: {TENANT_ID}" ^\n  -H "X-API-Key: {API_KEY}"')
add_body(doc, "若 connectivity-check 尚未初始化，得到 HTTP 404 是正确结果，表示域名、TLS、鉴权和数据库查询链路均已打通。得到 401 表示 Header 或密钥有误。")
add_heading(doc, "3.3 Windows PowerShell 一次性检查", 2)
add_code(doc, f'''$base = "{BASE_URL}"
$tenant = "{TENANT_ID}"
$key = "{API_KEY}"
$headers = @{{"X-Tenant-ID"=$tenant; "X-API-Key"=$key}}
Invoke-RestMethod "$base/health"
try {{ Invoke-RestMethod "$base/v1/profiles/connectivity-check" -Headers $headers }}
catch {{ $_.Exception.Response.StatusCode.value__ }}  # 首次应为 404''')

add_heading(doc, "4. HTTP 约定", 1)
add_table(
    doc,
    ["项目", "要求"],
    [
        ("传输", "HTTPS；JSON；UTF-8"),
        ("鉴权", "所有 /v1 请求带 X-Tenant-ID 与 X-API-Key"),
        ("写请求", "所有 POST 额外带 Idempotency-Key 与 Content-Type: application/json"),
        ("用户标识", "租户内稳定唯一且不可复用；不要直接用手机号、昵称或会话 ID"),
        ("版本控制", "写请求使用最近一次成功响应的 profile_version"),
        ("请求追踪", "记录响应头 X-Request-ID；报障时提供，禁止记录 API Key"),
        ("超时建议", "连接 5 秒、总请求 30 秒；网络错误可按幂等规则有限重试"),
    ],
    [2100, 7260],
)
add_heading(doc, "4.1 必带 Header", 2)
add_code(doc, f'''X-Tenant-ID: {TENANT_ID}
X-API-Key: {API_KEY}
Content-Type: application/json          # POST 请求
Idempotency-Key: <稳定且唯一的业务操作ID> # POST 请求''')
add_heading(doc, "4.2 幂等键规则", 2)
add_bullet(doc, "同一个业务操作因超时而重试时，Idempotency-Key 和请求体必须完全相同。")
add_bullet(doc, "不同接口、不同用户、不同消息不得复用同一个幂等键。")
add_bullet(doc, "推荐：初始化用 init:<USER_ID>；消息摄取直接使用机器人侧的 message_id。")
add_bullet(doc, "同一幂等键配不同请求体会返回 422；此时不要继续盲目重试。")

add_heading(doc, "5. 必须实现的主调用流程", 1)
add_callout(doc, "推荐顺序", "收到用户消息 → 读取或确认画像版本 → 必要时初始化 → 调用 messages:ingest → 把 reply_hints 注入机器人回答策略 → 保存新 profile_version → 调用聊天模型生成回复。", fill=COLORS["light_blue"], label_color=COLORS["navy"])

add_heading(doc, "5.1 读取已有画像", 2)
add_code(doc, f'''curl -sS "{BASE_URL}/v1/profiles/$USER_ID" \\
  -H "X-Tenant-ID: {TENANT_ID}" \\
  -H "X-API-Key: {API_KEY}"''')
add_bullet(doc, "HTTP 200：读取并保存 profile_version。")
add_bullet(doc, "HTTP 404：该用户还没有画像，执行 5.2 初始化。")

add_heading(doc, "5.2 首次初始化", 2)
add_code(doc, f'''curl -sS -X POST "{BASE_URL}/v1/profiles:init" \\
  -H "X-Tenant-ID: {TENANT_ID}" \\
  -H "X-API-Key: {API_KEY}" \\
  -H "Idempotency-Key: init:robot-user-001" \\
  -H "Content-Type: application/json" \\
  -d '{{
    "tenant_user_id": "robot-user-001",
    "display_name": "用户昵称",
    "consent": {{
      "profile": true,
      "sensitive_inference": false
    }}
  }}' ''')
add_body(doc, "初始化成功返回 profile_version=1。只有用户明确同意画像处理后，consent.profile 才能为 true。生日推断或九型数据属于敏感推断，只有取得相应授权后才能把 sensitive_inference 设为 true。")
add_body(doc, "可选初始化字段：birth_date（YYYY-MM-DD）、birth_time、timezone、enneagram。没有可靠信息时不要传空字符串或猜测值。")

add_heading(doc, "5.3 每轮摄取用户消息", 2)
add_code(doc, f'''curl -sS -X POST "{BASE_URL}/v1/profiles/robot-user-001/messages:ingest" \\
  -H "X-Tenant-ID: {TENANT_ID}" \\
  -H "X-API-Key: {API_KEY}" \\
  -H "Idempotency-Key: turn-001" \\
  -H "Content-Type: application/json" \\
  -d '{{
    "conversation_id": "session-001",
    "message_id": "turn-001",
    "expected_profile_version": 1,
    "occurred_at": "2026-08-07T08:00:00Z",
    "text": "以后回答短一点，先听我把话说完。",
    "context": {{
      "topic": "communication",
      "previous_turn_count": 0,
      "recent_turns": []
    }}
  }}' ''')
add_table(
    doc,
    ["字段", "要求"],
    [
        ("conversation_id", "本次会话稳定 ID，1-256 字符"),
        ("message_id", "当前用户消息唯一 ID，1-256 字符；推荐兼作幂等键"),
        ("expected_profile_version", "最近一次成功读取或写入得到的版本，最小为 1"),
        ("occurred_at", "带时区的 ISO 8601 时间，推荐 UTC"),
        ("text", "当前用户原话，1-10000 字符"),
        ("recent_turns", "可选最近上下文，最多 12 条；role 为 user/assistant；每条最多 4000 字符"),
    ],
    [2500, 6860],
)
add_body(doc, "成功响应中机器人必须消费或保存的字段：")
add_code(doc, '''{
  "profile_version": 2,
  "reply_hints": {
    "intent": "conversation",
    "tone": "natural",
    "max_sentences": 3,
    "answer_first": false,
    "empathy_first": true,
    "question_count": 1,
    "structure_level": "simple",
    "focus": "直接回应用户当前表达",
    "avoid": [],
    "requires_fresh_information": false
  },
  "profile_patch": [],
  "runtime_operations": [],
  "no_profile_change": false,
  "request_id": "req_xxx"
}''')
add_bullet(doc, "reply_hints：本轮机器人的回复长度、共情、回答顺序和结构建议。")
add_bullet(doc, "profile_version：保存为该用户最新版本，下一次写请求使用它。")
add_bullet(doc, "no_profile_change=true：本轮没有足够可靠的画像变化，不是错误；版本可能保持不变。")
add_bullet(doc, "profile_patch/runtime_operations：审计或高级集成使用，普通机器人回复链路可以不逐项解释。")
add_callout(doc, "时序选择", "如果画像策略要影响当前回答，messages:ingest 必须在调用聊天模型前同步执行，并把 reply_hints 写入系统提示词或回答参数。如果优先降低延迟，可以先用已有画像回答、再异步摄取；此时新策略从下一轮生效。", fill=COLORS["callout"], label_color=COLORS["dark_blue"])

add_heading(doc, "6. reply_hints 如何接到机器人回答", 1)
add_body(doc, "建议在机器人系统提示词中增加一个只由服务端生成的策略段，禁止让终端用户直接覆盖。示例：")
add_code(doc, '''【本轮回答策略】
- 先回答问题：{answer_first}
- 先表达共情：{empathy_first}
- 最多句数：{max_sentences}
- 追问数量：{question_count}
- 结构：{structure_level}
- 当前重点：{focus}
- 避免事项：{avoid}

规则：这些策略只控制表达方式，不改变事实，不替代安全策略和业务规则。''')
add_bullet(doc, "requires_fresh_information=true 时，先走机器人公司的联网搜索/实时数据流程，再生成事实性回答。")
add_bullet(doc, "avoid 为空时无需额外处理；非空时作为禁止表达清单。")
add_bullet(doc, "max_sentences 是表达约束，不应通过截断字符破坏句子或代码。")

add_heading(doc, "7. 可直接参考的 Node.js 18+ 最小实现", 1)
add_body(doc, "以下代码只依赖 Node.js 18+ 自带 fetch。生产中请把三个配置项放入密钥管理系统，不要硬编码。")
add_code(doc, '''const BASE_URL = process.env.PROFILE_ENGINE_BASE_URL;
const TENANT_ID = process.env.PROFILE_ENGINE_TENANT_ID;
const API_KEY = process.env.PROFILE_ENGINE_API_KEY;

async function api(path, { method = "GET", body, idemKey } = {}) {
  const headers = {
    "X-Tenant-ID": TENANT_ID,
    "X-API-Key": API_KEY,
  };
  if (body !== undefined) {
    headers["Content-Type"] = "application/json";
    headers["Idempotency-Key"] = idemKey;
  }
  const res = await fetch(`${BASE_URL}${path}`, {
    method, headers,
    body: body === undefined ? undefined : JSON.stringify(body),
    signal: AbortSignal.timeout(30000),
  });
  const data = await res.json().catch(() => ({}));
  if (!res.ok) {
    const err = new Error(`Profile API ${res.status}: ${data.code || data.detail}`);
    err.status = res.status;
    err.requestId = res.headers.get("x-request-id");
    err.body = data;
    throw err;
  }
  return data;
}

export async function ensureProfile(userId, displayName) {
  try {
    return await api(`/v1/profiles/${encodeURIComponent(userId)}`);
  } catch (err) {
    if (err.status !== 404) throw err;
    return api("/v1/profiles:init", {
      method: "POST",
      idemKey: `init:${userId}`,
      body: {
        tenant_user_id: userId,
        display_name: displayName,
        consent: { profile: true, sensitive_inference: false },
      },
    });
  }
}

export async function ingestTurn(userId, profileVersion, turn) {
  try {
    return await api(`/v1/profiles/${encodeURIComponent(userId)}/messages:ingest`, {
      method: "POST",
      idemKey: turn.messageId,
      body: {
        conversation_id: turn.conversationId,
        message_id: turn.messageId,
        expected_profile_version: profileVersion,
        occurred_at: turn.occurredAt,
        text: turn.text,
        context: { topic: turn.topic, recent_turns: turn.recentTurns || [] },
      },
    });
  } catch (err) {
    if (err.status !== 409) throw err;
    // 重新读取最新版本，用派生幂等键重试一次。
    const latest = await api(`/v1/profiles/${encodeURIComponent(userId)}`);
    return api(`/v1/profiles/${encodeURIComponent(userId)}/messages:ingest`, {
      method: "POST",
      idemKey: `${turn.messageId}:retry1`,
      body: {
        conversation_id: turn.conversationId,
        message_id: turn.messageId,
        expected_profile_version: latest.profile_version,
        occurred_at: turn.occurredAt,
        text: turn.text,
        context: { topic: turn.topic, recent_turns: turn.recentTurns || [] },
      },
    });
  }
}''')
add_callout(doc, "实现注意", "409 后请求体中的 expected_profile_version 会变化，因此若服务端将幂等键严格绑定完整请求体，重试应使用新的派生幂等键（例如 messageId + ':retry1'）。本服务当前以请求体摘要校验幂等键，故推荐重新读版本后使用派生键，避免 422。最多重试一次。", fill=COLORS["warning"], label_color=COLORS["warning_text"])

add_heading(doc, "8. 并发、409 与重试策略", 1)
add_number(doc, "每个用户串行处理画像写请求；不同用户可以并发。")
add_number(doc, "每次成功写入后，原子保存响应的 profile_version。")
add_number(doc, "收到 409 时，重新 GET 当前画像，取得最新 profile_version。")
add_number(doc, "用新的 expected_profile_version 和派生幂等键重试一次；仍冲突则进入队列或人工排查，禁止无限循环。")
add_number(doc, "网络超时且不知道服务是否已处理时，先用完全相同的幂等键和请求体重试；不要生成新键。")
add_table(
    doc,
    ["场景", "是否重试", "方式"],
    [
        ("连接超时/网络中断", "是", "同一幂等键 + 同一请求体，最多 2 次，指数退避"),
        ("409 版本冲突", "是", "重新 GET；新版本 + 派生幂等键，最多 1 次"),
        ("401/403/422", "否", "修正配置、授权或请求字段"),
        ("503", "是", "1s/2s/4s 指数退避，最多 3 次"),
        ("5xx 其他", "有限", "记录 X-Request-ID；最多 2 次后告警"),
    ],
    [2500, 1600, 5260],
    aligns=["left", "center", "left"],
)

add_heading(doc, "9. 可选管理接口", 1)
add_table(
    doc,
    ["用途", "接口"],
    [
        ("查看画像及证据依据", "GET /v1/profiles/{user_id}/explain?field=<字段路径>"),
        ("用户/专家明确更正", "POST /v1/profiles/{user_id}:correct"),
        ("写入已授权九型测评", "POST /v1/profiles/{user_id}:set-enneagram"),
        ("撤销记忆/证据/生日/九型/全部画像", "POST /v1/profiles/{user_id}:forget"),
        ("查询生产规则包", "GET /v1/rule-packs/current"),
    ],
    [3200, 6160],
)
add_heading(doc, "9.1 用户明确更正", 2)
add_code(doc, '''POST /v1/profiles/{user_id}:correct
{
  "expected_profile_version": 2,
  "target_path": "core_traits.energy_mode.extroversion",
  "value": 0.8,
  "reason": "用户明确更正"
}''')
add_heading(doc, "9.2 写入九型测评", 2)
add_code(doc, '''POST /v1/profiles/{user_id}:set-enneagram
{
  "expected_profile_version": 2,
  "enneagram": {
    "core_type": 5,
    "wing": 6,
    "primary_instinct": "SP",
    "secondary_instinct": "SO",
    "source": "external_assessment",
    "confidence": 0.9
  },
  "reason": "机器人公司已完成并取得授权的测评结果"
}''')
add_body(doc, "九型需要 sensitive_inference=true。wing 必须是主型相邻类型；两个 instinct 不能相同；core_type 为 1-9。")
add_heading(doc, "9.3 撤回或遗忘", 2)
add_code(doc, '''POST /v1/profiles/{user_id}:forget
{
  "expected_profile_version": 3,
  "scope": "all_profile",
  "reason": "用户撤回画像授权"
}''')
add_body(doc, "scope 可选 memory、evidence、birth_inference、enneagram、all_profile。memory/evidence 必须同时提供 target_id。all_profile 会停止后续画像推断并清理运行时偏好、状态、记忆、数字密码和九型数据，使证据/记忆失效；它不会物理删除历史审计记录。物理删除需走双方另行约定的数据删除流程。")
add_callout(doc, "禁止调用", "不要调用 POST /v1/profiles/{user_id}:reset。该接口只用于开发测试，生产环境固定返回 404；即使在线 OpenAPI 中可见，也不属于对外生产合同。", fill=COLORS["warning"], label_color=COLORS["red"])

add_heading(doc, "10. 错误码与现场处理", 1)
add_table(
    doc,
    ["HTTP", "含义", "负责人处理"],
    [
        ("400", "请求语法或业务格式错误", "检查 JSON、URL 与 Content-Type"),
        ("401", "租户或 API Key 错误", "停止重试；核对两个 Header 与密钥管理配置"),
        ("403", "未取得画像/敏感推断授权，或画像已关闭", "不自动重试；回到用户授权流程"),
        ("404", "用户不存在；或生产禁用的测试接口", "读取场景可初始化；reset 场景必须停止"),
        ("409", "profile_version 已过期", "重新读取最新版本后最多重试一次"),
        ("422", "Header/字段校验失败，或幂等键与请求体冲突", "修正请求；不要原样无限重试"),
        ("503", "语义处理临时不可用", "有限指数退避；保留 X-Request-ID"),
        ("5xx", "服务端异常", "有限重试并告警；提供时间、路径、状态、X-Request-ID"),
    ],
    [900, 2800, 5660],
    aligns=["center", "left", "left"],
)
add_body(doc, "标准错误体示例：")
add_code(doc, '''{
  "request_id": "req_xxx",
  "code": "profile_version_conflict",
  "message": "画像版本不匹配",
  "details": {
    "expected_profile_version": 1,
    "actual_profile_version": 2
  }
}''')
add_body(doc, "FastAPI Header/字段校验错误可能使用 detail 数组，而不是上述 code 结构。客户端应以 HTTP 状态为主，并兼容读取 code、detail、message。")

add_heading(doc, "11. 安全、隐私与数据边界", 1)
add_bullet(doc, "API 只能由机器人公司服务端调用；客户端、网页、小程序或机器人终端不得持有生产 Key。")
add_bullet(doc, "密钥应放入 Secret Manager/KMS 或受控环境变量；日志、告警、埋点、错误堆栈均需脱敏。")
add_bullet(doc, "tenant_user_id 使用内部不可逆 ID；不要直接传手机号、身份证、邮箱或昵称作为主键。")
add_bullet(doc, "只发送实现画像和回复策略所必需的消息上下文；recent_turns 保持最少必要。")
add_bullet(doc, "未取得 profile 授权不得初始化；未取得 sensitive_inference 授权不得提交生日推断或九型数据。")
add_bullet(doc, "当前生产语义提取器为本服务内置确定性规则，不把用户原话发送给外部语义模型。")
add_bullet(doc, "如密钥疑似泄露，立即停止使用并联系交付方轮换；不要把泄露的 Key 再粘贴到普通聊天或工单。")

add_heading(doc, "12. 上线验收清单", 1)
add_table(
    doc,
    ["验收项", "通过标准"],
    [
        ("HTTPS", "无需忽略证书校验即可访问正式域名"),
        ("健康检查", "HTTP 200；status=ok；application=ok；database=ok"),
        ("错误 Key", "返回 401"),
        ("租户隔离", "正确 Key + 错误租户返回 401"),
        ("初始化", "新用户返回 profile_version=1"),
        ("读取", "能按同一 tenant_user_id 读取画像"),
        ("消息摄取", "返回 reply_hints 与 profile_version"),
        ("幂等", "相同幂等键 + 相同请求体重复提交结果一致"),
        ("并发版本", "旧 profile_version 写入返回 409"),
        ("生产边界", "reset 返回 404"),
    ],
    [3300, 6060],
)
add_callout(doc, "我方已验收", "2026-08-07 已使用本文正式域名、租户与密钥完成线上 10/10 验收：健康、数据库、错误密钥、租户隔离、初始化、读取、消息摄取、幂等、409、生产禁用 reset 全部通过。", fill=COLORS["success"], label_color=COLORS["success_text"])

add_heading(doc, "13. 联调报障信息模板", 1)
add_body(doc, "出现问题时，只提供以下内容；不要提供 API Key、完整用户原话或完整画像：")
add_code(doc, '''环境：production
发生时间：YYYY-MM-DD HH:mm:ss（注明时区）
请求方法与路径：POST /v1/...
HTTP 状态：xxx
X-Request-ID：...
tenant_user_id：建议仅提供末 6 位或内部可追踪的脱敏值
是否为重试：是/否
期望结果：...
实际结果：...
已确认：健康检查结果 / 出站网络 / Header 是否存在''')

add_heading(doc, "14. 最终交接说明", 1)
add_body(doc, "机器人公司负责人完成以下五件事即可正式跑通：")
add_number(doc, "安全保存第 2 节的三个生产变量。")
add_number(doc, "完成第 3 节连通性验证。")
add_number(doc, "实现第 5 节主链路，并按用户维度保存 profile_version。")
add_number(doc, "把 reply_hints 注入本轮机器人回答策略。")
add_number(doc, "按第 12 节逐项验收；通过后上线。")
add_callout(doc, "唯一交付文件", "本指南已经包含运行地址、真实凭据、完整调用顺序、可复制命令、Node.js 示例、错误处理、安全边界和验收标准。接入负责人无需再索取其他文件即可开始联调。", fill=COLORS["light_blue"], label_color=COLORS["navy"])

# Avoid isolated headings at page bottoms and reduce awkward paragraph splits.
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
    elif len(paragraph.text) < 800:
        paragraph.paragraph_format.widow_control = True

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(str(OUTPUT))
